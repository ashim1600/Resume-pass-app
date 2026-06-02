"""Resume Builder App

Streamlit app that takes a job description and a base profile/resume, then
generates a tailored resume draft.

Run:
    pip install -r requirements.txt
    streamlit run app.py
"""

from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass
from io import BytesIO
from html import escape
from typing import Iterable

import streamlit as st

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None

try:
    import docx
    from docx.shared import Pt
except Exception:  # pragma: no cover
    docx = None
    Pt = None


st.set_page_config(page_title="Resume Builder", page_icon="📄", layout="wide", initial_sidebar_state="expanded")


def inject_styles() -> None:
    st.markdown(
        """
        <style>
        :root {
            --bg: #f4f7fb;
            --panel: rgba(255, 255, 255, 0.88);
            --panel-border: rgba(15, 23, 42, 0.08);
            --ink: #0f172a;
            --muted: #475569;
            --accent: #2563eb;
            --accent-2: #7c3aed;
            --accent-soft: rgba(37, 99, 235, 0.09);
            --shadow: 0 18px 50px rgba(15, 23, 42, 0.08);
            --radius: 18px;
        }

        .stApp {
            background:
                radial-gradient(circle at top left, rgba(37, 99, 235, 0.10), transparent 30%),
                radial-gradient(circle at top right, rgba(124, 58, 237, 0.08), transparent 28%),
                linear-gradient(180deg, #f8fbff 0%, #f4f7fb 48%, #eef2ff 100%);
            color: var(--ink);
        }

        .stApp,
        .stApp p,
        .stApp span,
        .stApp label,
        .stApp div,
        .stApp li,
        .stApp a,
        .stApp h1,
        .stApp h2,
        .stApp h3,
        .stApp h4,
        .stApp h5,
        .stApp h6 {
            color: var(--ink);
        }

        [data-testid="stHeader"] {
            background: transparent;
        }

        [data-testid="stToolbar"] {
            right: 12px;
        }

        .block-container {
            padding-top: 2rem;
            padding-bottom: 2.25rem;
            max-width: 1420px;
        }

        .hero {
            background: linear-gradient(135deg, rgba(15, 23, 42, 0.96), rgba(37, 99, 235, 0.95) 45%, rgba(124, 58, 237, 0.92));
            color: white;
            border-radius: 28px;
            padding: 2rem 2rem 1.7rem 2rem;
            box-shadow: var(--shadow);
            border: 1px solid rgba(255,255,255,0.12);
        }

        .hero h1 {
            margin: 0;
            font-size: 2.35rem;
            line-height: 1.05;
            letter-spacing: 0;
        }

        .hero p {
            margin: 0.75rem 0 0 0;
            max-width: 72ch;
            font-size: 1.02rem;
            line-height: 1.6;
            color: rgba(255, 255, 255, 0.88);
        }

        .badge-row {
            display: flex;
            flex-wrap: wrap;
            gap: 0.5rem;
            margin-top: 1rem;
        }

        .badge {
            display: inline-flex;
            align-items: center;
            gap: 0.45rem;
            padding: 0.42rem 0.75rem;
            border-radius: 999px;
            background: rgba(255,255,255,0.12);
            border: 1px solid rgba(255,255,255,0.14);
            color: white;
            font-size: 0.82rem;
            line-height: 1;
        }

        .section-title {
            font-size: 1rem;
            font-weight: 700;
            letter-spacing: 0;
            color: var(--ink);
            margin: 0 0 0.65rem 0;
        }

        .glass-card {
            background: var(--panel);
            backdrop-filter: blur(16px);
            border: 1px solid var(--panel-border);
            border-radius: var(--radius);
            box-shadow: var(--shadow);
            padding: 1rem 1rem 0.9rem 1rem;
        }

        .metric-card {
            background: var(--panel);
            border: 1px solid var(--panel-border);
            border-radius: 16px;
            box-shadow: var(--shadow);
            padding: 1rem 1rem 0.85rem 1rem;
        }

        .metric-label {
            font-size: 0.78rem;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: var(--muted);
            margin-bottom: 0.25rem;
        }

        .metric-value {
            font-size: 1.5rem;
            font-weight: 800;
            color: var(--ink);
            line-height: 1.1;
        }

        .metric-subtext {
            margin-top: 0.3rem;
            color: var(--muted);
            font-size: 0.88rem;
        }

        .resume-paper {
            background: #fff;
            color: #111827;
            border: 1px solid rgba(15, 23, 42, 0.08);
            border-radius: 20px;
            box-shadow: 0 18px 60px rgba(15, 23, 42, 0.10);
            padding: 1.5rem 1.55rem;
        }

        .resume-name {
            font-size: 1.75rem;
            font-weight: 800;
            letter-spacing: 0;
            margin: 0;
            color: #0f172a;
        }

        .resume-title {
            margin-top: 0.15rem;
            font-size: 1rem;
            color: #334155;
            font-weight: 600;
        }

        .resume-contact {
            margin-top: 0.55rem;
            color: #475569;
            font-size: 0.92rem;
            line-height: 1.5;
        }

        .resume-divider {
            height: 1px;
            background: linear-gradient(90deg, rgba(37, 99, 235, 0.35), rgba(124, 58, 237, 0.15), transparent);
            margin: 1rem 0;
        }

        .resume-section {
            margin-bottom: 1rem;
        }

        .resume-section h3 {
            margin: 0 0 0.45rem 0;
            font-size: 0.95rem;
            letter-spacing: 0.12em;
            text-transform: uppercase;
            color: #0f172a;
        }

        .resume-section p,
        .resume-section li {
            margin: 0;
            color: #1e293b;
            line-height: 1.55;
            font-size: 0.95rem;
        }

        .resume-section ul {
            margin: 0.15rem 0 0 1.05rem;
            padding: 0;
        }

        .chip-grid {
            display: flex;
            flex-wrap: wrap;
            gap: 0.45rem;
        }

        .chip {
            display: inline-flex;
            align-items: center;
            padding: 0.42rem 0.68rem;
            border-radius: 999px;
            background: rgba(37, 99, 235, 0.08);
            color: #0f172a;
            border: 1px solid rgba(37, 99, 235, 0.12);
            font-size: 0.86rem;
        }

        .hint-box {
            background: rgba(255,255,255,0.72);
            border: 1px solid rgba(15, 23, 42, 0.08);
            border-radius: 16px;
            padding: 0.9rem 1rem;
            color: #334155;
        }

        .stTabs [data-baseweb="tab-list"] {
            gap: 0.35rem;
            background: rgba(255,255,255,0.68);
            padding: 0.35rem;
            border-radius: 14px;
            border: 1px solid rgba(15, 23, 42, 0.06);
        }

        .stTabs [data-baseweb="tab"] {
            border-radius: 11px;
            padding: 0.65rem 0.95rem;
            font-weight: 600;
        }

        .stTabs [aria-selected="true"] {
            background: white;
            box-shadow: 0 8px 24px rgba(15,23,42,0.08);
        }

        .stTextArea textarea, .stTextInput input {
            border-radius: 14px !important;
            border-color: rgba(15, 23, 42, 0.10) !important;
            background: rgba(255, 255, 255, 0.92) !important;
            color: #0f172a !important;
            -webkit-text-fill-color: #0f172a !important;
        }

        .stTextArea textarea::placeholder,
        .stTextInput input::placeholder {
            color: #64748b !important;
            opacity: 1 !important;
        }

        [data-testid="stFileUploader"] *,
        [data-testid="stSidebar"] *,
        [data-testid="stForm"] *,
        [data-testid="stTabs"] * {
            color: #0f172a;
        }

        [data-testid="stMetric"] *,
        [data-testid="stMetricValue"],
        [data-testid="stMetricLabel"] {
            color: #0f172a !important;
        }

        .stMarkdown,
        .stMarkdown p,
        .stMarkdown li,
        .stMarkdown span {
            color: #0f172a;
        }

        .stInfo,
        .stWarning,
        .stError,
        .stSuccess {
            color: #0f172a !important;
        }

        div[data-testid="stFileUploader"] {
            background: rgba(255,255,255,0.75);
            border-radius: 16px;
            border: 1px dashed rgba(37, 99, 235, 0.18);
            padding: 0.2rem 0.4rem 0.6rem 0.4rem;
        }

        .stButton > button {
            border-radius: 14px;
            border: 0;
            background: linear-gradient(135deg, #2563eb, #7c3aed);
            color: white;
            font-weight: 700;
            padding: 0.72rem 1rem;
            box-shadow: 0 14px 28px rgba(37, 99, 235, 0.24);
        }

        .stDownloadButton > button {
            border-radius: 14px !important;
            border: 1px solid rgba(15, 23, 42, 0.08) !important;
            background: white !important;
            color: #0f172a !important;
            font-weight: 700 !important;
        }

        .callout-grid {
            display: grid;
            grid-template-columns: repeat(3, minmax(0, 1fr));
            gap: 0.75rem;
        }

        .callout {
            background: rgba(255,255,255,0.86);
            border: 1px solid rgba(15, 23, 42, 0.08);
            border-radius: 16px;
            padding: 0.9rem 1rem;
            box-shadow: 0 12px 28px rgba(15, 23, 42, 0.05);
        }

        .callout .kicker {
            font-size: 0.77rem;
            text-transform: uppercase;
            letter-spacing: 0.08em;
            color: #64748b;
            margin-bottom: 0.3rem;
            font-weight: 700;
        }

        .callout .title {
            font-size: 0.98rem;
            font-weight: 800;
            color: #0f172a;
            margin-bottom: 0.25rem;
        }

        .callout .body {
            color: #334155;
            font-size: 0.9rem;
            line-height: 1.45;
        }

        .section-shell {
            background: rgba(255,255,255,0.76);
            border: 1px solid rgba(15, 23, 42, 0.08);
            border-radius: 20px;
            padding: 1rem;
            box-shadow: var(--shadow);
        }

        .subtle-label {
            text-transform: uppercase;
            letter-spacing: 0.10em;
            font-size: 0.75rem;
            color: #64748b;
            font-weight: 700;
            margin-bottom: 0.45rem;
        }

        .summary-box {
            border-radius: 18px;
            border: 1px solid rgba(15, 23, 42, 0.08);
            background: white;
            padding: 1rem 1rem 0.9rem 1rem;
            box-shadow: 0 10px 24px rgba(15, 23, 42, 0.05);
        }

        .summary-box h4 {
            margin: 0 0 0.45rem 0;
            font-size: 0.95rem;
            font-weight: 800;
            color: #0f172a;
        }

        .summary-box p {
            margin: 0;
            color: #334155;
            line-height: 1.5;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


STOPWORDS = {
    "a",
    "about",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "has",
    "have",
    "he",
    "her",
    "his",
    "i",
    "in",
    "into",
    "is",
    "it",
    "its",
    "me",
    "of",
    "on",
    "or",
    "our",
    "she",
    "that",
    "the",
    "their",
    "them",
    "this",
    "to",
    "was",
    "we",
    "were",
    "with",
    "you",
    "your",
}


SECTION_PATTERNS = {
    "experience": re.compile(r"(?im)^\s*(experience|work experience|professional experience)\s*$"),
    "education": re.compile(r"(?im)^\s*(education|academic background)\s*$"),
    "skills": re.compile(r"(?im)^\s*(skills|technical skills|core skills)\s*$"),
    "projects": re.compile(r"(?im)^\s*(projects|selected projects)\s*$"),
    "summary": re.compile(r"(?im)^\s*(summary|professional summary|profile)\s*$"),
}


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def tokenize(text: str) -> list[str]:
    words = re.findall(r"[a-zA-Z][a-zA-Z0-9+#.\-/]+", text.lower())
    return [w for w in words if w not in STOPWORDS and len(w) > 1]


def extract_keywords(text: str, max_keywords: int = 40) -> list[str]:
    tokens = tokenize(text)
    if not tokens:
        return []
    counts = Counter(tokens)
    return [word for word, _ in counts.most_common(max_keywords)]


def extract_file_text(uploaded_file) -> str:
    if uploaded_file is None:
        return ""

    raw = uploaded_file.getvalue()
    name = (uploaded_file.name or "").lower()

    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")

    if name.endswith(".pdf"):
        if PdfReader is None:
            return ""
        reader = PdfReader(BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if name.endswith(".docx"):
        if docx is None:
            return ""
        buffer = BytesIO(raw)
        document = docx.Document(buffer)
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    return ""


def split_lines(text: str) -> list[str]:
    return [line.strip("-• \t") for line in (text or "").splitlines() if line.strip("-• \t").strip()]


def sentence_snippets(text: str, keywords: Iterable[str], limit: int = 8) -> list[str]:
    sentences = re.split(r"(?<=[.!?])\s+|\n+", text)
    hits: list[str] = []
    for sentence in sentences:
        s_norm = normalize(sentence)
        if not s_norm:
            continue
        if any(keyword in s_norm for keyword in keywords):
            hits.append(sentence.strip())
    return hits[:limit]


def section_presence(text: str) -> dict[str, bool]:
    return {name: bool(pattern.search(text)) for name, pattern in SECTION_PATTERNS.items()}


def infer_title(job_text: str, fallback: str) -> str:
    first_line = next((line.strip() for line in (job_text or "").splitlines() if line.strip()), "")
    if first_line and len(first_line) < 80:
        cleaned = re.sub(r"[^a-zA-Z0-9 &/\-+]", "", first_line)
        if cleaned:
            return cleaned
    return fallback or "Target Role"


def extract_name_from_resume(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines() if line.strip()]
    return lines[0] if lines else ""


def extract_contact_block(text: str) -> dict[str, str]:
    raw = text or ""
    email = next(iter(re.findall(r"[\w.\-+%]+@[\w.\-]+\.[A-Za-z]{2,}", raw)), "")
    phone = next(iter(re.findall(r"\+?\d[\d\s().-]{7,}\d", raw)), "")
    linkedin = next(iter(re.findall(r"https?://(?:www\.)?linkedin\.com/[^\s)]+", raw, flags=re.I)), "")
    github = next(iter(re.findall(r"https?://(?:www\.)?github\.com/[^\s)]+", raw, flags=re.I)), "")
    location = ""
    for line in split_lines(raw):
        if any(token in line.lower() for token in ("remote", "india", "united states", "usa", "hyderabad", "bangalore", "bengaluru", "mumbai", "delhi", "pune", "chennai")):
            location = line
            break
    return {
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "location": location,
    }


def rank_resume_lines(resume_text: str, keywords: Iterable[str]) -> list[str]:
    lines = split_lines(resume_text)
    scored: list[tuple[int, str]] = []
    for line in lines:
        line_norm = normalize(line)
        score = sum(1 for kw in keywords if kw in line_norm)
        if score:
            scored.append((score, line))
    scored.sort(key=lambda x: (-x[0], len(x[1])))
    return [line for _, line in scored]


def build_summary(name: str, title: str, keywords: list[str], existing_summary: str) -> str:
    top = ", ".join(keywords[:6]) if keywords else "relevant tools and technologies"
    if existing_summary.strip():
        base = existing_summary.strip().rstrip(".")
    else:
        base = f"Results-driven {title}"
    return f"{base} with strengths in {top}. Focused on delivering measurable outcomes, clear communication, and practical execution."


def build_skills_section(job_keywords: list[str], base_resume_text: str) -> list[str]:
    base_keywords = extract_keywords(base_resume_text, max_keywords=20)
    combined: list[str] = []
    seen = set()
    for keyword in job_keywords + base_keywords:
        if keyword not in seen:
            seen.add(keyword)
            combined.append(keyword)
    return combined[:18]


def build_experience_bullets(job_keywords: list[str], base_resume_text: str, limit: int = 6) -> list[str]:
    ranked_lines = rank_resume_lines(base_resume_text, job_keywords)
    bullets: list[str] = []
    for line in ranked_lines:
        cleaned = line.strip()
        if len(cleaned) < 20:
            continue
        bullets.append(cleaned)
        if len(bullets) >= limit:
            break

    if bullets:
        return bullets

    return [
        f"Delivered work aligned with {', '.join(job_keywords[:4]) if job_keywords else 'the target role'}.",
        "Collaborated with stakeholders to clarify requirements, priorities, and delivery timelines.",
        "Improved processes by standardizing recurring tasks and documenting repeatable workflows.",
    ]


def build_projects(job_keywords: list[str], base_resume_text: str) -> list[str]:
    project_lines = []
    for line in split_lines(base_resume_text):
        if len(line) > 24 and any(k in normalize(line) for k in job_keywords[:10]):
            project_lines.append(line)
    if project_lines:
        return project_lines[:3]
    if job_keywords:
        return [f"Created a project or initiative using {', '.join(job_keywords[:3])}."]
    return ["Added an optional project here that matches the role."]


def build_education(base_resume_text: str, education_input: str) -> list[str]:
    if education_input.strip():
        return split_lines(education_input)
    lines = split_lines(base_resume_text)
    edu_lines = [line for line in lines if any(word in normalize(line) for word in ("university", "college", "school", "bachelor", "master", "mba", "btech", "mtech", "degree", "diploma"))]
    return edu_lines[:3] or ["Add your education details here."]


def parse_skill_items(text: str) -> list[str]:
    return [item.strip() for item in re.split(r"[,;\n]+", text or "") if item.strip()]


def section_to_lines(text: str, limit: int = 8) -> list[str]:
    items = []
    for line in split_lines(text):
        cleaned = line.strip()
        if len(cleaned) > 14:
            items.append(cleaned)
        if len(items) >= limit:
            break
    return items


def fit_breakdown(job_keywords: list[str], resume_text: str, sections: dict[str, list[str]]) -> dict[str, int]:
    resume_norm = normalize(resume_text)
    keyword_fit = round((sum(1 for kw in job_keywords if kw in resume_norm) / max(len(job_keywords), 1)) * 100)
    section_fit = 0
    section_fit += 20 if sections.get("summary") else 0
    section_fit += 20 if sections.get("skills") else 0
    section_fit += 20 if sections.get("experience") else 0
    section_fit += 20 if sections.get("projects") else 0
    section_fit += 20 if sections.get("education") else 0
    contact_fit = 0
    if re.search(r"[\w.\-+%]+@[\w.\-]+\.[A-Za-z]{2,}", resume_text):
        contact_fit += 40
    if re.search(r"\+?\d[\d\s().-]{7,}\d", resume_text):
        contact_fit += 30
    if re.search(r"https?://|linkedin\.com|github\.com", resume_norm):
        contact_fit += 30
    source_depth = min(100, round(len(tokenize(resume_text)) / 5))
    return {
        "keyword_fit": keyword_fit,
        "section_fit": section_fit,
        "contact_fit": contact_fit,
        "source_depth": source_depth,
    }


def build_action_items(job_keywords: list[str], missing_keywords: list[str], checks: list[str], resume_text: str) -> list[str]:
    items: list[str] = []
    if missing_keywords:
        items.append(f"Add or mirror the top missing terms: {', '.join(missing_keywords[:5])}.")
    if checks:
        items.append(checks[0])
    if len(tokenize(resume_text)) < 180:
        items.append("Add one more quantified accomplishment to strengthen the draft.")
    if not items:
        items.append("The source is adequate. Refine one or two bullets for sharper impact.")
    return items[:3]


def resume_tone_summary(tone: str) -> str:
    mapping = {
        "Direct": "Concise, objective, and ATS-first",
        "Balanced": "Polished, modern, and broadly professional",
        "Confident": "Sharper, stronger, and more executive",
    }
    return mapping.get(tone, "Polished, modern, and broadly professional")


def apply_tone(text: str, tone: str) -> str:
    if tone == "Direct":
        return text.replace("Focused on delivering measurable outcomes, clear communication, and practical execution.", "Focused on execution and measurable outcomes.")
    if tone == "Confident":
        return text.replace("Results-driven", "Highly capable").replace("Focused on delivering", "Known for delivering")
    return text


def assemble_result(
    *,
    profile: dict[str, str],
    summary: str,
    skills: list[str],
    experience: list[str],
    projects: list[str],
    education: list[str],
    job_keywords: list[str],
    resume_final: str,
    tone: str,
    template: str,
    job_final: str,
) -> dict[str, object]:
    resume_text_output = build_resume_text(
        profile,
        {
            "summary": [summary],
            "skills": skills,
            "experience": experience,
            "projects": projects,
            "education": education,
        },
        template=template,
    )
    resume_text_output = apply_tone(resume_text_output, tone)
    resume_norm = normalize(resume_text_output)
    matched_keywords = [kw for kw in job_keywords if kw in resume_norm]
    missing_keywords = [kw for kw in job_keywords if kw not in resume_norm]
    score = round((len(matched_keywords) / len(job_keywords)) * 100) if job_keywords else 0
    sections = {
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "projects": projects,
        "education": education,
    }
    fit = fit_breakdown(job_keywords, resume_text_output, sections)
    checks = quality_checks(resume_final)
    actions = build_action_items(job_keywords, missing_keywords, checks, resume_final)
    fit["overall"] = round((fit["keyword_fit"] * 0.45) + (fit["section_fit"] * 0.25) + (fit["contact_fit"] * 0.15) + (fit["source_depth"] * 0.15))
    fit["overall"] = min(100, max(score, fit["overall"]))
    fit["overall"] = round(fit["overall"])
    return {
        "profile": profile,
        "summary": summary,
        "skills": skills,
        "experience": experience,
        "projects": projects,
        "education": education,
        "resume_text_output": resume_text_output,
        "score": score,
        "matched_keywords": matched_keywords,
        "missing_keywords": missing_keywords,
        "job_keywords": job_keywords,
        "resume_final": resume_final,
        "job_final": job_final,
        "tone": tone,
        "template": template,
        "fit": fit,
        "actions": actions,
    }


def build_resume_text(profile: dict[str, str], sections: dict[str, list[str]], template: str = "Modern") -> str:
    contact_items = [profile["email"], profile["phone"], profile["location"], profile["linkedin"], profile["github"]]
    contact = " | ".join(item for item in contact_items if item)

    template = (template or "Modern").lower()
    if template == "classic":
        order = ["summary", "experience", "skills", "projects", "education"]
    elif template == "executive":
        order = ["summary", "skills", "experience", "projects", "education"]
    else:
        order = ["summary", "skills", "experience", "projects", "education"]

    lines = []
    lines.append(profile["name"] or "Your Name")
    lines.append(profile["title"] or "Target Title")
    if contact:
        lines.append(contact)
    for section_name in order:
        lines.append("")
        lines.append(section_name.upper())
        if section_name == "summary":
            lines.append(sections["summary"][0])
        elif section_name == "skills":
            lines.append(", ".join(sections["skills"]))
        elif section_name == "experience":
            lines.extend(f"- {bullet}" for bullet in sections["experience"])
        elif section_name == "projects":
            lines.extend(f"- {bullet}" for bullet in sections["projects"])
        elif section_name == "education":
            lines.extend(f"- {item}" for item in sections["education"])
    return "\n".join(lines).strip() + "\n"


def build_resume_docx(text: str) -> bytes:
    if docx is None:
        return b""

    document = docx.Document()
    if Pt is not None:
        style = document.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)

    for block in text.strip().split("\n\n"):
        lines = [line for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        if len(lines) == 1 and lines[0].isupper():
            document.add_paragraph(lines[0]).style = "Heading 1"
            continue
        if lines[0].isupper() and len(lines) > 1:
            document.add_paragraph(lines[0]).style = "Heading 1"
            for line in lines[1:]:
                if line.startswith("- "):
                    document.add_paragraph(line[2:], style="List Bullet")
                else:
                    document.add_paragraph(line)
            continue
        for line in lines:
            if line.startswith("- "):
                document.add_paragraph(line[2:], style="List Bullet")
            else:
                document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_metric_card(label: str, value: str, subtext: str) -> str:
    return f"""
    <div class="metric-card">
      <div class="metric-label">{escape(label)}</div>
      <div class="metric-value">{escape(value)}</div>
      <div class="metric-subtext">{escape(subtext)}</div>
    </div>
    """


def render_chip_row(items: Iterable[str]) -> str:
    chips = "".join(f'<span class="chip">{escape(item)}</span>' for item in items if item)
    return f'<div class="chip-grid">{chips}</div>' if chips else '<div class="chip-grid"></div>'


def render_resume_preview(profile: dict[str, str], sections: dict[str, list[str]]) -> str:
    contact_bits = [profile["email"], profile["phone"], profile["location"], profile["linkedin"], profile["github"]]
    contact = " | ".join(escape(bit) for bit in contact_bits if bit)
    summary = escape(sections["summary"][0]) if sections["summary"] else ""
    skills = "".join(f"<span class='chip'>{escape(item)}</span>" for item in sections["skills"])
    exp_items = "".join(f"<li>{escape(item)}</li>" for item in sections["experience"])
    project_items = "".join(f"<li>{escape(item)}</li>" for item in sections["projects"])
    edu_items = "".join(f"<li>{escape(item)}</li>" for item in sections["education"])

    return f"""
    <div class="resume-paper">
      <div class="resume-name">{escape(profile["name"])}</div>
      <div class="resume-title">{escape(profile["title"])}</div>
      <div class="resume-contact">{contact}</div>
      <div class="resume-divider"></div>

      <div class="resume-section">
        <h3>Summary</h3>
        <p>{summary}</p>
      </div>

      <div class="resume-section">
        <h3>Skills</h3>
        <div class="chip-grid">{skills}</div>
      </div>

      <div class="resume-section">
        <h3>Experience</h3>
        <ul>{exp_items}</ul>
      </div>

      <div class="resume-section">
        <h3>Projects</h3>
        <ul>{project_items}</ul>
      </div>

      <div class="resume-section">
        <h3>Education</h3>
        <ul>{edu_items}</ul>
      </div>
    </div>
    """


def analyze_job(job_text: str, max_keywords: int) -> list[str]:
    return extract_keywords(job_text, max_keywords=max_keywords)


def quality_checks(resume_text: str) -> list[str]:
    checks: list[str] = []
    text = normalize(resume_text)
    if len(text.split()) < 120:
        checks.append("Resume source is short. Add more experience details for a stronger draft.")
    if "@" not in text:
        checks.append("No email address detected in the source material.")
    if not re.search(r"\b(\+?\d[\d\s().-]{7,}\d)\b", resume_text):
        checks.append("No phone number detected in the source material.")
    sections = section_presence(resume_text)
    if not sections["experience"]:
        checks.append("No Experience heading detected in the source material.")
    if not sections["skills"]:
        checks.append("No Skills heading detected in the source material.")
    if not sections["education"]:
        checks.append("No Education heading detected in the source material.")
    return checks


def main() -> None:
    inject_styles()
    st.session_state.setdefault("resume_builder_result", None)
    st.session_state.setdefault("resume_versions", [])

    st.title("Resume Builder")
    st.markdown(
        """
        <div class="hero">
          <h1>Build a tailored resume from a job description.</h1>
          <p>
            Upload or paste the job description, add your real background details, and generate a polished
            resume draft with a cleaner structure, stronger keyword alignment, and downloadable output.
          </p>
          <div class="badge-row">
            <span class="badge">ATS-aware draft</span>
            <span class="badge">DOCX and TXT export</span>
            <span class="badge">Designed preview</span>
            <span class="badge">Fact-based output</span>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="callout-grid">
          <div class="callout">
            <div class="kicker">Workflow</div>
            <div class="title">Paste or upload</div>
            <div class="body">Bring in the job description and your real background, then generate a structured draft.</div>
          </div>
          <div class="callout">
            <div class="kicker">Output</div>
            <div class="title">Editable resume</div>
            <div class="body">Refine the generated sections before you download the final TXT or DOCX version.</div>
          </div>
          <div class="callout">
            <div class="kicker">Quality</div>
            <div class="title">Fit analysis</div>
            <div class="body">See keyword coverage, section completeness, and concrete next actions to improve the draft.</div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        st.header("Generation Controls")
        max_keywords = st.slider("Job keywords to use", 10, 60, 30, 5)
        bullet_limit = st.slider("Experience bullets", 3, 10, 6)
        tone = st.select_slider("Writing tone", options=["Direct", "Balanced", "Confident"], value="Balanced")
        template = st.radio("Template", ["Modern", "Classic", "Executive"], index=0)
        output_length = st.select_slider("Output length", options=["Concise", "Standard", "Detailed"], value="Standard")
        st.divider()
        st.markdown(
            '<div class="hint-box">The draft only uses the facts you give it. Better input produces a cleaner resume.</div>',
            unsafe_allow_html=True,
        )

    with st.form("resume_builder_form"):
        left, right = st.columns(2, gap="large")

        with left:
            st.markdown('<div class="section-title">Job Description</div>', unsafe_allow_html=True)
            jd_file = st.file_uploader("Upload job description", type=["pdf", "docx", "txt"], key="jd_file")
            jd_text = st.text_area("Or paste the job description", height=300, placeholder="Paste the job description here...")
            target_role = st.text_input("Target role", value="Software Engineer")

        with right:
            st.markdown('<div class="section-title">Your Base Profile</div>', unsafe_allow_html=True)
            resume_file = st.file_uploader("Upload your existing resume", type=["pdf", "docx", "txt"], key="resume_file")
            resume_text = st.text_area("Or paste your existing resume / profile notes", height=300, placeholder="Paste your current resume or notes here...")

        c1, c2, c3 = st.columns(3)
        with c1:
            full_name = st.text_input("Full name", value="")
            email = st.text_input("Email", value="")
        with c2:
            phone = st.text_input("Phone", value="")
            location = st.text_input("Location", value="")
        with c3:
            linkedin = st.text_input("LinkedIn", value="")
            github = st.text_input("GitHub", value="")

        summary_input = st.text_area("Optional summary to keep", height=110, placeholder="Paste a summary you already want to keep, or leave blank.")
        skills_input = st.text_area("Optional skills to keep", height=110, placeholder="One skill per line or comma-separated.")
        education_input = st.text_area("Education details", height=110, placeholder="Paste your education details here.")

        submit = st.form_submit_button("Generate tailored resume", use_container_width=True)

    if submit:
        jd_source = extract_file_text(jd_file)
        resume_source = extract_file_text(resume_file)
        job_final = (jd_text.strip() or jd_source.strip()).strip()
        resume_final = (resume_text.strip() or resume_source.strip()).strip()

        if jd_file and not jd_source and not jd_text.strip():
            st.warning("Could not extract text from the uploaded job description. Paste the text directly if needed.")
        if resume_file and not resume_source and not resume_text.strip():
            st.warning("Could not extract text from the uploaded resume. Paste the text directly if needed.")

        if not job_final:
            st.error("Add a job description first.")
            st.session_state["resume_builder_result"] = None
            return
        if not (resume_final or summary_input.strip() or skills_input.strip() or education_input.strip()):
            st.error("Add at least some source material about your background before generating.")
            st.session_state["resume_builder_result"] = None
            return

        job_keywords = analyze_job(job_final, max_keywords=max_keywords)
        existing_summary = summary_input.strip()
        extracted_name = extract_name_from_resume(resume_final)
        extracted_contact = extract_contact_block(resume_final)

        profile = {
            "name": full_name.strip() or extracted_name or "Your Name",
            "title": target_role.strip() or infer_title(job_final, "Target Role"),
            "email": email.strip() or extracted_contact["email"],
            "phone": phone.strip() or extracted_contact["phone"],
            "location": location.strip() or extracted_contact["location"],
            "linkedin": linkedin.strip() or extracted_contact["linkedin"],
            "github": github.strip() or extracted_contact["github"],
        }

        summary = build_summary(profile["name"], profile["title"], job_keywords, existing_summary)
        skills = build_skills_section(job_keywords, skills_input or resume_final)
        if skills_input.strip():
            skills = parse_skill_items(skills_input)[:18] or skills

        if output_length == "Concise":
            exp_limit = max(3, bullet_limit - 2)
            proj_limit = 1
        elif output_length == "Detailed":
            exp_limit = min(10, bullet_limit + 2)
            proj_limit = 3
        else:
            exp_limit = bullet_limit
            proj_limit = 2

        experience = build_experience_bullets(job_keywords, resume_final, limit=exp_limit)
        projects = build_projects(job_keywords, resume_final)[:proj_limit]
        education = build_education(resume_final, education_input)

        summary = apply_tone(summary, tone)

        st.session_state["resume_builder_result"] = assemble_result(
            profile=profile,
            summary=summary,
            skills=skills,
            experience=experience,
            projects=projects,
            education=education,
            job_keywords=job_keywords,
            resume_final=resume_final,
            tone=tone,
            template=template,
            job_final=job_final,
        )
        st.session_state["resume_versions"].append(st.session_state["resume_builder_result"])

    result = st.session_state.get("resume_builder_result")
    if not result:
        st.info("Fill the inputs, then generate the resume draft.")
        return

    tabs = st.tabs(["Overview", "Preview", "Edit", "Download"])
    profile = result["profile"]
    summary = result["summary"]
    skills = result["skills"]
    experience = result["experience"]
    projects = result["projects"]
    education = result["education"]
    resume_text_output = result["resume_text_output"]
    score = result["score"]
    matched_keywords = result["matched_keywords"]
    missing_keywords = result["missing_keywords"]
    job_keywords = result["job_keywords"]
    resume_final = result["resume_final"]
    fit = result["fit"]
    actions = result["actions"]
    tone = result["tone"]
    template = result["template"]
    section_presence_map = section_presence(result["resume_text_output"])

    with tabs[0]:
        metric_cols = st.columns(4)
        metric_payloads = [
            ("Overall fit", f"{fit['overall']}%", "Blend of keyword, section, contact, and source strength"),
            ("Matched keywords", f"{len(matched_keywords)}/{len(job_keywords)}", "From the job description"),
            ("Template", template, "Controls layout order"),
            ("Tone", resume_tone_summary(tone), "Controls voice of the draft"),
        ]
        for col, payload in zip(metric_cols, metric_payloads):
            with col:
                st.markdown(render_metric_card(*payload), unsafe_allow_html=True)

        overview_cols = st.columns([1.1, 0.9], gap="large")
        with overview_cols[0]:
            st.markdown('<div class="subtle-label">Fit Breakdown</div>', unsafe_allow_html=True)
            break_cols = st.columns(4)
            for col, (label, value) in zip(
                break_cols,
                [
                    ("Keyword", f"{fit['keyword_fit']}%"),
                    ("Section", f"{fit['section_fit']}%"),
                    ("Contact", f"{fit['contact_fit']}%"),
                    ("Depth", f"{fit['source_depth']}%"),
                ],
            ):
                with col:
                    st.markdown(render_metric_card(label, value, "Score component"), unsafe_allow_html=True)

            st.markdown('<div class="subtle-label" style="margin-top:1rem;">Next Actions</div>', unsafe_allow_html=True)
            for item in actions:
                st.markdown(f'<div class="summary-box" style="margin-bottom:0.6rem;"><h4>Action</h4><p>{escape(item)}</p></div>', unsafe_allow_html=True)

        with overview_cols[1]:
            st.markdown('<div class="subtle-label">Section Completeness</div>', unsafe_allow_html=True)
            completeness = {
                "Summary": "Present" if section_presence_map["summary"] else "Missing",
                "Skills": "Present" if section_presence_map["skills"] else "Missing",
                "Experience": "Present" if section_presence_map["experience"] else "Missing",
                "Projects": "Present" if section_presence_map["projects"] else "Missing",
                "Education": "Present" if section_presence_map["education"] else "Missing",
            }
            for name, value in completeness.items():
                st.markdown(f'<div class="hint-box" style="margin-bottom:0.55rem;"><strong>{escape(name)}:</strong> {escape(value)}</div>', unsafe_allow_html=True)

            st.markdown('<div class="subtle-label" style="margin-top:1rem;">Job Keywords</div>', unsafe_allow_html=True)
            st.markdown(render_chip_row(job_keywords[:18]), unsafe_allow_html=True)

    with tabs[1]:
        st.markdown(render_resume_preview(profile, {
            "summary": [summary],
            "skills": skills,
            "experience": experience,
            "projects": projects,
            "education": education,
        }), unsafe_allow_html=True)

    with tabs[2]:
        st.markdown("### Refine the generated resume")
        st.markdown("Edit the generated content before export. These changes update the download output immediately.")

        with st.form("refine_form"):
            p1, p2 = st.columns(2)
            with p1:
                edit_name = st.text_input("Name", value=profile["name"])
                edit_title = st.text_input("Target title", value=profile["title"])
                edit_email = st.text_input("Email", value=profile["email"])
                edit_phone = st.text_input("Phone", value=profile["phone"])
            with p2:
                edit_location = st.text_input("Location", value=profile["location"])
                edit_linkedin = st.text_input("LinkedIn", value=profile["linkedin"])
                edit_github = st.text_input("GitHub", value=profile["github"])
                edit_summary = st.text_area("Summary", value=summary, height=120)

            skill_text = st.text_area("Skills", value=", ".join(skills), height=110)
            experience_text = st.text_area("Experience bullets", value="\n".join(f"- {item}" for item in experience), height=180)
            projects_text = st.text_area("Projects", value="\n".join(f"- {item}" for item in projects), height=140)
            education_text = st.text_area("Education", value="\n".join(f"- {item}" for item in education), height=120)

            save_edits = st.form_submit_button("Apply edits", use_container_width=True)

        if save_edits:
            new_profile = {
                "name": edit_name.strip() or "Your Name",
                "title": edit_title.strip() or "Target Title",
                "email": edit_email.strip(),
                "phone": edit_phone.strip(),
                "location": edit_location.strip(),
                "linkedin": edit_linkedin.strip(),
                "github": edit_github.strip(),
            }
            updated_result = assemble_result(
                profile=new_profile,
                summary=edit_summary.strip() or summary,
                skills=parse_skill_items(skill_text) or skills,
                experience=section_to_lines(experience_text) or experience,
                projects=section_to_lines(projects_text) or projects,
                education=section_to_lines(education_text) or education,
                job_keywords=job_keywords,
                resume_final=resume_final,
                tone=tone,
                template=template,
                job_final=result["job_final"],
            )
            st.session_state["resume_builder_result"] = updated_result
            st.session_state["resume_versions"].append(updated_result)
            st.success("Edits applied. The preview and downloads now use the updated draft.")
            st.rerun()

    with tabs[3]:
        st.download_button(
            "Download as TXT",
            data=resume_text_output.encode("utf-8"),
            file_name=f"{profile['name'].strip().replace(' ', '_') or 'resume'}_tailored.txt",
            mime="text/plain",
            use_container_width=True,
        )

        docx_bytes = build_resume_docx(resume_text_output)
        if docx_bytes:
            st.download_button(
                "Download as DOCX",
                data=docx_bytes,
                file_name=f"{profile['name'].strip().replace(' ', '_') or 'resume'}_tailored.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        else:
            st.info("Install python-docx to enable DOCX download.")

        st.markdown("### Export text")
        st.code(resume_text_output, language="text")

        if st.session_state["resume_versions"]:
            st.markdown("### Draft history")
            history_labels = [f"Version {idx + 1}" for idx in range(len(st.session_state["resume_versions"]))]
            st.selectbox("Saved versions", history_labels, index=len(history_labels) - 1, key="resume_history_select")


if __name__ == "__main__":
    main()
